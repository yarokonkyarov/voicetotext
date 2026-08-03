#!/usr/bin/env python3
"""
Voice-to-Text Telegram Bot
Transcribes MP3/WAV audio via Deepgram and formats using PLAUD-style templates.
"""

import asyncio
import hashlib
import hmac
import json
import os
import logging
import tempfile
import subprocess
from datetime import datetime, timedelta, timezone
import httpx
from typing import Dict, Any, Optional

from aiohttp import web
from docx import Document
from docx.shared import Pt
from dotenv import load_dotenv
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup, Message, constants
from telegram.ext import (
    Application,
    MessageHandler,
    CallbackQueryHandler,
    CommandHandler,
    filters,
    ContextTypes,
)
from deepgram import DeepgramClient, DeepgramClientOptions, PrerecordedOptions, FileSource
import anthropic
from pyrogram import Client as PyroClient
from todoist_integration import (
    TODOIST_ENABLED_TEMPLATES,
    TODOIST_PROMPT_SUFFIX,
    extract_action_items,
    build_todoist_button,
    todoist_callback,
)

load_dotenv()

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
)
logger = logging.getLogger(__name__)

TELEGRAM_BOT_TOKEN = os.getenv("TELEGRAM_BOT_TOKEN", "8385003837:AAGcZz-LP5exdsFy4i5mn578RRey-6mJcRM")
DEEPGRAM_API_KEY = os.getenv("DEEPGRAM_API_KEY", "8b296b68246aaa4e5468512f972a4d9ea90dfb8a")
ANTHROPIC_API_KEY = os.getenv("ANTHROPIC_API_KEY", "")
TELEGRAM_API_ID = int(os.getenv("TELEGRAM_API_ID", "0"))
TELEGRAM_API_HASH = os.getenv("TELEGRAM_API_HASH", "")
SECOND_BRAIN_URL = os.getenv("SECOND_BRAIN_URL", "")
SECOND_BRAIN_API_KEY = os.getenv("SECOND_BRAIN_API_KEY", "")
# plaud-processor: текст/txt/md уходят туда (общий бот, поллер у plaud выключен)
PLAUD_PROCESSOR_URL = os.getenv("PLAUD_PROCESSOR_URL", "")
PLAUD_WEBHOOK_SECRET = os.getenv("PLAUD_WEBHOOK_SECRET", "")

# applaud (self-hosted Plaud-девайс синк) шлёт сюда transcript_ready вебхуки —
# запись готова, минуя ручную отправку файла в Telegram
MY_CHAT_ID = os.getenv("MY_CHAT_ID", "")
APPLAUD_WEBHOOK_SECRET = os.getenv("APPLAUD_WEBHOOK_SECRET", "")
APPLAUD_WEBHOOK_PORT = int(os.getenv("APPLAUD_WEBHOOK_PORT", "8020"))
MSK = timezone(timedelta(hours=3))

# ─── State ────────────────────────────────────────────────────────────────────

# user_id -> {path, filename}
pending_audio: Dict[str, Dict[str, Any]] = {}

# user_id -> {mode: "initial"|"reformat", path?: str, transcript?: str, chat_id: int}
pending_spec_context: Dict[str, Dict[str, Any]] = {}

# user_id -> {formatted, template_name, transcript}  — last processed result
last_result: Dict[str, Dict[str, Any]] = {}



# ─── Pyrogram client (large file downloads) ───────────────────────────────────

_pyro: Optional[PyroClient] = None


async def get_pyro() -> PyroClient:
    global _pyro
    if _pyro is None:
        _pyro = PyroClient(
            "voicetotext_dl",
            api_id=TELEGRAM_API_ID,
            api_hash=TELEGRAM_API_HASH,
            bot_token=TELEGRAM_BOT_TOKEN,
            workdir="/root/voicetotext",
            no_updates=True,
        )
        await _pyro.start()
    return _pyro


async def download_large_file(chat_id: int, message_id: int, suffix: str) -> str:
    client = await get_pyro()
    msg = await client.get_messages(chat_id, message_id)
    tmp = tempfile.NamedTemporaryFile(suffix=suffix, delete=False)
    tmp.close()
    await client.download_media(msg, file_name=tmp.name)
    return tmp.name

# ─── Templates ────────────────────────────────────────────────────────────────

TEMPLATES: Dict[str, Dict[str, str]] = {
    "meeting": {
        "name": "📋 Meeting Notes",
        "description": "Ключевые пункты, решения и задачи",
    },
    "interview": {
        "name": "🎤 Interview",
        "description": "Формат вопрос-ответ с выводами",
    },
    "lecture": {
        "name": "📚 Lecture Notes",
        "description": "Концепции, определения и конспект",
    },
    "journal": {
        "name": "📔 Voice Journal",
        "description": "Личные заметки, мысли и планы",
    },
    "summary": {
        "name": "✨ Smart Summary",
        "description": "TL;DR и ключевые тезисы",
    },
    "spec": {
        "name": "📐 Тех. задание",
        "description": "Функциональные требования к продукту",
    },
    "client_call": {
        "name": "📞 Звонок с клиентом",
        "description": "Потребности, возражения, договорённости, следующие шаги",
    },
    "email": {
        "name": "✉️ Письмо",
        "description": "Голос → готовое письмо с темой и структурой",
    },
    "sales_coaching": {
        "name": "🎯 Оценка звонка",
        "description": "Коучинг по 5 этапам: знакомство, квалификация, презентация, возражения, закрытие",
    },
    "client_meeting": {
        "name": "📋 Client Meeting Notes",
        "description": "Notitie met klantbehoefte, producten, voorwaarden en vervolgacties",
    },
}

TEMPLATE_PROMPTS: Dict[str, str] = {
    "meeting": """Ты профессиональный редактор протоколов встреч. Оформи транскрипт ниже в структурированный протокол на русском языке.

## 📋 Протокол встречи

### Краткое резюме
[2–3 предложения о чём была встреча]

### Ключевые темы обсуждения
- ...

### Принятые решения
- ...

### Задачи и поручения
- [ ] Задача — Ответственный (если упомянут)

### Следующие шаги
- ...

---
Транскрипт:
{transcript}

Верни только оформленный протокол. Никаких пояснений.""",

    "interview": """Ты профессиональный редактор интервью. Оформи транскрипт ниже в формат интервью на русском языке.

## 🎤 Интервью

[Раздели на вопросы и ответы. Если спикеры не обозначены — определи самостоятельно.]

**В:** ...
**О:** ...

### Ключевые выводы
- ...

---
Транскрипт:
{transcript}

Верни только оформленное интервью. Никаких пояснений.""",

    "lecture": """Ты редактор учебных материалов. Оформи транскрипт ниже как конспект лекции на русском языке.

## 📚 Конспект лекции

### Тема
[Основной предмет]

### Ключевые концепции
1. ...

### Определения и термины
**Термин:** Определение

### Краткое содержание
[Краткий абзац]

### Вопросы для повторения
- ...

---
Транскрипт:
{transcript}

Верни только конспект. Никаких пояснений.""",

    "journal": """Ты помогаешь вести голосовой дневник. Оформи транскрипт ниже как запись в дневнике на русском языке.

## 📔 Запись в дневнике

### Мысли и размышления
[Основные мысли из записи]

### Важные моменты
[Значимые события или идеи]

### Настроение и эмоции
[Эмоциональный контекст, если прослеживается]

### Планы и намерения
[Упомянутые цели или задачи]

---
Транскрипт:
{transcript}

Верни только запись в дневнике. Никаких пояснений.""",

    "summary": """Ты эксперт по созданию кратких резюме. Создай умное резюме транскрипта на русском языке.

## ✨ Умное резюме

### Коротко о главном (TL;DR)
[2–3 предложения]

### Ключевые тезисы
• ...
• ...
• ...

### Важные детали
[Факты и подробности, которые стоит запомнить]

### Что нужно сделать
[Конкретные действия и следующие шаги, если упомянуты]

---
Транскрипт:
{transcript}

Верни только резюме. Никаких пояснений.""",

    "client_call": """Ты эксперт по продажам и работе с клиентами. Оформи транскрипт ниже как структурированную заметку о звонке с клиентом на русском языке.

## 📞 Звонок с клиентом

### Клиент и контекст
[Кто звонил / с кем говорили, компания, должность если упомянуты]

### Потребности и запросы клиента
- ...

### Возражения и опасения
- ...

### Достигнутые договорённости
- ...

### Следующие шаги
- [ ] Задача — Ответственный (если упомянут) — Срок (если упомянут)

### Общее впечатление
[Настрой клиента, уровень заинтересованности, риски]

---
Транскрипт:
{transcript}

Верни только заметку о звонке. Никаких пояснений.""",

    "sales_coaching": """You are an expert sales coach. Analyze the transcript below and produce a structured call evaluation report in the same language as the transcript.

## 🎯 Sales Call Evaluation

### 1. Introduction
**Score:** ⭐⭐⭐⭐⭐ (X/5)
[How did the agent open the call? First impression, tone, clarity of purpose]

### 2. Qualifying Questions
**Score:** ⭐⭐⭐⭐⭐ (X/5)
[Did the agent ask the right questions to understand customer needs and build rapport?]

### 3. Presentation
**Score:** ⭐⭐⭐⭐⭐ (X/5)
[How well did the agent present the product/service, key features, and promotions?]

### 4. Objection Handling & Sales Tactics
**Score:** ⭐⭐⭐⭐⭐ (X/5)
[How did the agent handle objections? What sales techniques were used?]

### 5. Closing
**Score:** ⭐⭐⭐⭐⭐ (X/5)
[Did the agent attempt to close? Was the closing technique effective?]

---

### ✅ Strengths
- ...

### 📈 Areas for Improvement
- ...

### 💬 Overall Coaching Summary
[2–3 sentences on overall performance and key focus for next call]

---
Transcript:
{transcript}

Return only the evaluation report. No extra commentary.""",

    "client_meeting": """Ты опытный аккаунт-менеджер. Составь структурированную заметку о встрече с клиентом на основе транскрипта ниже.

## 📋 Встреча с клиентом — [название компании] — [дата]

### Потребность клиента
[Чего хочет достичь клиент? Какую проблему нужно решить?]

### Обсуждаемые продукты / услуги
- ...

### Технические требования и условия
- ...

### Возражения и важные моменты
- ...

### Договорённости и следующие шаги
- [ ] Действие — Ответственный — Срок

---
Транскрипт:
{transcript}

Верни только заметку. Никаких пояснений.""",

    "email": """Ты профессиональный бизнес-копирайтер. Преобразуй транскрипт ниже в готовое деловое письмо на русском языке.

## ✉️ Письмо

**Тема:** [Сформулируй тему письма]

---

[Приветствие]

[Основная часть письма — структурированно, по существу, деловой тон]

[Завершение и призыв к действию если нужен]

С уважением,
[Имя автора если упомянуто]

---
Транскрипт:
{transcript}

Верни только готовое письмо. Никаких пояснений.""",

    "spec": """Ты эксперт по написанию технических заданий и функциональных требований. Оформи транскрипт ниже как структурированное ТЗ на русском языке.

## 📐 Функциональные требования

### Описание продукта
[Что это за продукт, для кого предназначен, какую проблему решает]

### Цели и задачи
- ...

### Функциональные требования
#### FR-1. [Название функциональности]
- **Описание:** ...
- **Входные данные:** ...
- **Ожидаемый результат:** ...

#### FR-2. ...

### Нефункциональные требования
- **Производительность:** ...
- **Безопасность:** ...
- **Интерфейс и UX:** ...
- **Надёжность:** ...

### Ограничения и допущения
- ...

### Открытые вопросы
- ...

---
Транскрипт:
{transcript}

Верни только ТЗ. Никаких пояснений.""",
}

SPEC_CUSTOM_PROMPT = """Ты эксперт по написанию технических заданий и функциональных требований.

Контекст от заказчика:
{custom_context}

Используя этот контекст как направляющий, оформи транскрипт ниже как структурированное ТЗ на русском языке.

## 📐 Функциональные требования

### Описание продукта
[Что это за продукт, для кого предназначен, какую проблему решает]

### Цели и задачи
- ...

### Функциональные требования
#### FR-1. [Название функциональности]
- **Описание:** ...
- **Входные данные:** ...
- **Ожидаемый результат:** ...

#### FR-2. ...

### Нефункциональные требования
- **Производительность:** ...
- **Безопасность:** ...
- **Интерфейс и UX:** ...
- **Надёжность:** ...

### Ограничения и допущения
- ...

### Открытые вопросы
- ...

---
Транскрипт:
{transcript}

Верни только ТЗ. Никаких пояснений."""

# ─── Keyboards ────────────────────────────────────────────────────────────────


def build_template_keyboard(user_id: str, prefix: str = "tpl") -> InlineKeyboardMarkup:
    keyboard = [
        [InlineKeyboardButton(
            f"{t['name']}  —  {t['description']}",
            callback_data=f"{prefix}:{tid}:{user_id}",
        )]
        for tid, t in TEMPLATES.items()
    ]
    if prefix == "retpl":
        keyboard.append([InlineKeyboardButton("✖ Отмена", callback_data=f"action:reformat_cancel:{user_id}")])
    return InlineKeyboardMarkup(keyboard)


def build_action_keyboard(user_id: str) -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup([
        [InlineKeyboardButton("🔄 Другой шаблон", callback_data=f"action:reformat:{user_id}")],
        [
            InlineKeyboardButton("📄 DOCX", callback_data=f"action:docx:{user_id}"),
            InlineKeyboardButton("📝 Markdown", callback_data=f"action:md:{user_id}"),
        ],
    ])


# ─── Core helpers ─────────────────────────────────────────────────────────────


async def send_long_message(chat_id: int, text: str, bot) -> Message:
    """Split and send messages longer than Telegram's 4096-char limit. Returns last message."""
    limit = 4000
    chunks = [text[i: i + limit] for i in range(0, len(text), limit)]
    last = None
    for chunk in chunks:
        last = await bot.send_message(chat_id=chat_id, text=chunk)
    return last



def extract_audio(input_path: str) -> str:
    out = tempfile.NamedTemporaryFile(suffix=".ogg", delete=False)
    out.close()
    subprocess.run(
        [
            "ffmpeg", "-y", "-i", input_path,
            "-vn",
            "-acodec", "libopus",
            "-b:a", "32k",
            "-ar", "16000",
            out.name,
        ],
        check=True,
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
    )
    return out.name


def transcribe_audio(tmp_path: str) -> str:
    deepgram = DeepgramClient(
        DEEPGRAM_API_KEY,
    )
    with open(tmp_path, "rb") as f:
        buffer_data = f.read()
    payload: FileSource = {"buffer": buffer_data}
    options = PrerecordedOptions(
        model="nova-2",
        smart_format=True,
        language="ru",
        punctuate=True,
        paragraphs=True,
        diarize=True,
    )
    response = deepgram.listen.rest.v("1").transcribe_file(payload, options)
    return response.results.channels[0].alternatives[0].transcript


def format_with_claude(template_id: str, transcript: str, custom_context: Optional[str] = None) -> str:
    client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)

    MAX_TRANSCRIPT_CHARS = 60_000
    if len(transcript) > MAX_TRANSCRIPT_CHARS:
        transcript = transcript[:MAX_TRANSCRIPT_CHARS] + "\n\n[... транскрипт обрезан ...]"

    if template_id == "spec" and custom_context:
        prompt = SPEC_CUSTOM_PROMPT.format(custom_context=custom_context.strip(), transcript=transcript)
    else:
        prompt = TEMPLATE_PROMPTS[template_id].format(transcript=transcript)

    if template_id in TODOIST_ENABLED_TEMPLATES:
        prompt += TODOIST_PROMPT_SUFFIX

    response = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=8096,
        messages=[{"role": "user", "content": prompt}],
    )
    raw_text = response.content[0].text
    clean_text, tasks = extract_action_items(raw_text)
    # tasks передаётся наружу через атрибут функции (простой способ без рефакторинга)
    format_with_claude._last_tasks = tasks
    return clean_text


def create_docx(template_name: str, formatted_text: str, source_filename: str = "") -> str:
    """Convert formatted text to .docx. Returns temp file path."""
    doc = Document()

    doc_title = f"{template_name} — {source_filename}" if source_filename else template_name
    title = doc.add_heading(doc_title, level=0)
    title.runs[0].font.size = Pt(18)

    for line in formatted_text.splitlines():
        if line.startswith("#### "):
            doc.add_heading(line[5:], level=3)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith(("- ", "• ")):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.startswith("- [ ] "):
            doc.add_paragraph("☐ " + line[6:], style="List Bullet")
        elif line.strip() in ("---", "─" * 10):
            doc.add_paragraph("─" * 40)
        elif line.strip():
            # Bold **text** → strip markers for simplicity
            clean = line.replace("**", "")
            doc.add_paragraph(clean)

    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    doc.save(tmp.name)
    tmp.close()
    return tmp.name



# ─── Result delivery ──────────────────────────────────────────────────────────


async def send_to_second_brain(
    template_id: str,
    template_name: str,
    transcript: str,
    formatted: str,
    source_filename: str,
    tasks: list,
):
    """Отправляет результат в second-brain (Obsidian vault). Ошибки не критичны."""
    if not (SECOND_BRAIN_URL and SECOND_BRAIN_API_KEY):
        return
    clean_name = template_name.lstrip("📋📞✉️🎯📐🔖📝 ")
    payload = {
        "source": "voicetotext",
        "title": f"{clean_name} — {source_filename}" if source_filename else clean_name,
        "text": formatted,
        "raw_text": transcript,
        "tags": ["voicetotext"] + TEMPLATE_TAGS.get(template_id, []),
        "tasks": [
            t["text"] + (f" (@{t['owner']})" if t.get("owner") else "")
            for t in tasks
            if t.get("text")
        ],
    }
    try:
        async with httpx.AsyncClient(timeout=60) as client:
            resp = await client.post(
                f"{SECOND_BRAIN_URL}/ingest",
                json=payload,
                headers={"X-API-Key": SECOND_BRAIN_API_KEY},
            )
            resp.raise_for_status()
            logger.info("second-brain: сохранено в %s", resp.json().get("note_path"))
    except Exception as e:
        logger.error("second-brain ingest failed: %s", e)


async def forward_to_plaud(transcript: str, source: str) -> bool:
    """Пересылает расшифровку в plaud-processor (пайплайн задач). True — принято."""
    if not PLAUD_PROCESSOR_URL:
        return False
    payload = {
        "transcript": transcript,
        "recording_id": f"tg-{source}-{datetime.now().strftime('%Y%m%d%H%M%S')}",
        "timestamp": datetime.now().astimezone().isoformat(),
    }
    headers = {"X-Webhook-Secret": PLAUD_WEBHOOK_SECRET} if PLAUD_WEBHOOK_SECRET else {}
    try:
        async with httpx.AsyncClient(timeout=120) as client:
            resp = await client.post(
                f"{PLAUD_PROCESSOR_URL}/webhook/plaud", json=payload, headers=headers,
            )
            resp.raise_for_status()
            return True
    except Exception as e:
        logger.error("plaud-processor forward failed: %s", e)
        return False


async def deliver_result(
    chat_id: int,
    bot,
    user_id: str,
    template_id: str,
    transcript: str,
    formatted: str,
    source_filename: str = "",
    tasks: list = None,
):
    """Send formatted result and store it with action buttons."""
    template_name = TEMPLATES[template_id]["name"]
    tasks = tasks or []

    asyncio.create_task(
        send_to_second_brain(
            template_id, template_name, transcript, formatted, source_filename, tasks
        )
    )

    last_result[user_id] = {
        "transcript": transcript,
        "formatted": formatted,
        "template_id": template_id,
        "template_name": template_name,
        "source_filename": source_filename,
    }

    header_parts = [f"🎙 {template_name}"]
    if source_filename:
        header_parts.append(f"📎 {source_filename}")
    header_parts.append("─" * 32)
    header = "\n".join(header_parts) + "\n\n"
    await send_long_message(chat_id=chat_id, text=header + formatted, bot=bot)

    # Строим клавиатуру с опциональной кнопкой Todoist
    kb_rows = [
        [InlineKeyboardButton("🔄 Другой шаблон", callback_data=f"action:reformat:{user_id}")],
        [
            InlineKeyboardButton("📄 DOCX", callback_data=f"action:docx:{user_id}"),
            InlineKeyboardButton("📝 Markdown", callback_data=f"action:md:{user_id}"),
        ],
    ]
    todoist_btn = build_todoist_button(tasks, source=source_filename)
    if todoist_btn:
        kb_rows.append([todoist_btn])

    await bot.send_message(
        chat_id=chat_id,
        text="Действия с результатом:",
        reply_markup=InlineKeyboardMarkup(kb_rows),
    )




# Tag mapping per template
TEMPLATE_TAGS = {
    "meeting":        ["встреча", "протокол"],
    "action_items":   ["задачи", "action-items"],
    "summary":        ["резюме", "summary"],
    "spec":           ["тех-задание", "spec"],
    "client_call":    ["клиент", "звонок", "crm"],
    "email":          ["письмо", "email"],
    "sales_coaching": ["продажи", "коучинг", "оценка-звонка"],
    "client_meeting": ["клиент", "встреча", "crm"],
}


def create_md(template_id: str, template_name: str, formatted_text: str, source_filename: str = "") -> str:
    """Generate Obsidian-compatible Markdown file. Returns temp file path."""
    today = datetime.now().strftime("%Y-%m-%d")
    clean_name = template_name.lstrip("📋📞✉️🎯📐🔖📝 ")  # strip emoji for tags

    # Build YAML frontmatter
    tags = ["voicetotext"] + TEMPLATE_TAGS.get(template_id, [])
    tags_yaml = "\n".join(f"  - {t}" for t in tags)

    title = f"{clean_name} — {source_filename}" if source_filename else clean_name

    frontmatter = (
        f"---\n"
        f"title: \"{title}\"\n"
        f"date: {today}\n"
        f"tags:\n{tags_yaml}\n"
        f"source: \"{source_filename}\"\n"
        f"template: \"{clean_name}\"\n"
        f"---\n\n"
    )

    md_content = frontmatter + formatted_text

    tmp = tempfile.NamedTemporaryFile(suffix=".md", delete=False, mode="w", encoding="utf-8")
    tmp.write(md_content)
    tmp.close()
    return tmp.name

# ─── Spec context flow ────────────────────────────────────────────────────────


async def process_after_spec_context(
    chat_id: int,
    status_message: Message,
    bot,
    user_id: str,
    spec_state: Dict[str, Any],
    custom_context: Optional[str],
):
    """Transcribe (if needed) and format as spec with optional custom context."""
    loop = asyncio.get_event_loop()
    try:
        if spec_state["mode"] == "initial":
            tmp_path = spec_state["path"]
            await status_message.edit_text(
                f"✅ Шаблон: {TEMPLATES['spec']['name']}\n⏳ Транскрибирую аудио…"
            )
            try:
                compressed_path = await loop.run_in_executor(None, extract_audio, tmp_path)
            except Exception:
                compressed_path = tmp_path
            try:
                transcript = await loop.run_in_executor(None, transcribe_audio, compressed_path)
            finally:
                for p in set([tmp_path, compressed_path]):
                    try:
                        os.unlink(p)
                    except OSError:
                        pass
        else:
            transcript = spec_state["transcript"]

        if not transcript.strip():
            await status_message.edit_text("❌ Deepgram не смог распознать речь. Попробуй снова.")
            return

        await status_message.edit_text(
            f"✅ Транскрипция готова ({len(transcript)} символов)\n⏳ Формирую ТЗ…"
        )

        if ANTHROPIC_API_KEY:
            try:
                formatted = await loop.run_in_executor(
                    None, format_with_claude, "spec", transcript, custom_context
                )
            except Exception as e:
                logger.exception("Claude failed")
                await bot.send_message(chat_id=chat_id, text=f"⚠️ Claude ошибка: {e}\n\nСырой транскрипт:")
                formatted = transcript
        else:
            formatted = transcript

        await status_message.edit_text(f"✅ Готово! Шаблон: {TEMPLATES['spec']['name']}")
        await deliver_result(chat_id, bot, user_id, "spec", transcript, formatted)

    except Exception as exc:
        logger.exception("Spec processing error")
        await status_message.edit_text(f"❌ Ошибка: {exc}")


# ─── Handlers ────────────────────────────────────────────────────────────────


async def cmd_start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text(
        "👋 Привет! Отправь мне аудиофайл (MP3, WAV, OGG, WEBM) или голосовое сообщение, "
        "и я транскрибирую его с помощью Deepgram.\n\n"
        "После загрузки выбери шаблон оформления текста."
    )


async def handle_audio(update: Update, context: ContextTypes.DEFAULT_TYPE):
    message = update.message
    user_id = str(update.effective_user.id)

    SUPPORTED_MIME = {
        "audio/mpeg", "audio/mp3",
        "audio/wav", "audio/wave", "audio/x-wav",
        "audio/ogg", "audio/opus",
        "audio/webm", "video/webm",
    }

    # Audio received while waiting for spec context voice → treat as context voice
    if user_id in pending_spec_context and (message.voice or message.audio):
        tg_file = await (message.voice or message.audio).get_file()
        suffix = ".ogg" if message.voice else ".mp3"
        tmp = tempfile.NamedTemporaryFile(suffix=suffix, delete=False)
        await tg_file.download_to_drive(tmp.name)
        tmp.close()

        status = await message.reply_text("⏳ Транскрибирую контекст…")
        loop = asyncio.get_event_loop()
        try:
            ctx_text = await loop.run_in_executor(None, transcribe_audio, tmp.name)
        except Exception as exc:
            await status.edit_text(f"❌ Не удалось транскрибировать контекст: {exc}")
            return
        finally:
            try:
                os.unlink(tmp.name)
            except OSError:
                pass

        spec_state = pending_spec_context.pop(user_id)
        await status.edit_text(f"✅ Контекст получен\n⏳ Обрабатываю запись…")
        await process_after_spec_context(
            chat_id=message.chat_id,
            status_message=status,
            bot=context.bot,
            user_id=user_id,
            spec_state=spec_state,
            custom_context=ctx_text,
        )
        return

    # Normal audio upload
    MAX_FILE_SIZE = 20 * 1024 * 1024  # 20 MB — Telegram bot limit

    if message.audio:
        media = message.audio
        filename = media.file_name or "audio.mp3"
    elif message.voice:
        media = message.voice
        filename = "voice.ogg"
    elif message.video:
        media = message.video
        filename = media.file_name or "video.webm"
    elif message.document and message.document.mime_type in SUPPORTED_MIME:
        media = message.document
        filename = media.file_name or "audio.mp3"
    else:
        await message.reply_text("⚠️ Пожалуйста, отправь аудиофайл MP3, WAV, OGG, WEBM или голосовое сообщение.")
        return

    if media.file_size and media.file_size > MAX_FILE_SIZE:
        size_mb = media.file_size / 1024 / 1024
        status = await message.reply_text(
            f"📦 Файл {size_mb:.1f} МБ — скачиваю через расширенный режим…"
        )
        suffix = os.path.splitext(filename)[-1] or ".webm"
        try:
            tmp_path = await download_large_file(message.chat_id, message.message_id, suffix)
        except Exception as exc:
            logger.exception("Pyrogram download failed")
            await status.edit_text(f"❌ Не удалось скачать файл: {exc}")
            return
        await status.edit_text("✅ Файл получен! Выбери шаблон оформления:")
        pending_audio[user_id] = {"path": tmp_path, "filename": filename}
        await status.edit_text(
            "✅ Файл получен! Выбери шаблон оформления:",
            reply_markup=build_template_keyboard(user_id),
        )
        return

    tg_file = await media.get_file()

    status = await message.reply_text("⏳ Загружаю файл…")
    suffix = os.path.splitext(filename)[-1] or ".mp3"
    tmp = tempfile.NamedTemporaryFile(suffix=suffix, delete=False)
    await tg_file.download_to_drive(tmp.name)
    tmp.close()

    pending_audio[user_id] = {"path": tmp.name, "filename": filename}
    await status.edit_text(
        "✅ Файл получен! Выбери шаблон оформления:",
        reply_markup=build_template_keyboard(user_id),
    )


async def handle_text_message(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handle text: spec context input, иначе — транскрипт для plaud-processor."""
    user_id = str(update.effective_user.id)
    text = update.message.text.strip()

    if user_id in pending_spec_context:
        spec_state = pending_spec_context.pop(user_id)
        status = await update.message.reply_text("✅ Контекст получен\n⏳ Обрабатываю запись…")
        await process_after_spec_context(
            chat_id=update.message.chat_id,
            status_message=status,
            bot=context.bot,
            user_id=user_id,
            spec_state=spec_state,
            custom_context=text,
        )
        return

    # Свободный текст = готовая расшифровка (Plaud) → пайплайн задач
    if PLAUD_PROCESSOR_URL:
        status = await update.message.reply_text("⏳ Разбираю расшифровку на задачи…")
        ok = await forward_to_plaud(text, str(update.message.message_id))
        if ok:
            await status.delete()
        else:
            await status.edit_text("❌ Не удалось обработать текст (plaud-processor недоступен).")


async def handle_transcript_file(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Файл .txt/.md = готовая расшифровка (Plaud) → пайплайн задач."""
    doc = update.message.document
    if doc.file_size and doc.file_size > 5 * 1024 * 1024:
        await update.message.reply_text("❌ Файл слишком большой (лимит 5 МБ).")
        return
    status = await update.message.reply_text("⏳ Разбираю расшифровку на задачи…")
    try:
        tg_file = await doc.get_file()
        raw = bytes(await tg_file.download_as_bytearray())
        text = raw.decode("utf-8", errors="replace")
    except Exception as e:
        logger.error("transcript file download failed: %s", e)
        await status.edit_text(f"❌ Не удалось скачать файл: {e}")
        return
    caption = (update.message.caption or "").strip()
    full_text = f"{caption}\n\n{text}" if caption else text
    ok = await forward_to_plaud(full_text, doc.file_name or str(update.message.message_id))
    if ok:
        await status.delete()
    else:
        await status.edit_text("❌ Не удалось обработать файл (plaud-processor недоступен).")


async def handle_template_selection(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Initial template selection after audio upload."""
    query = update.callback_query
    await query.answer()

    parts = query.data.split(":", 2)
    if parts[0] != "tpl" or len(parts) != 3:
        return
    _, template_id, user_id = parts

    if user_id not in pending_audio:
        await query.edit_message_text("❌ Аудиофайл не найден. Пожалуйста, отправь файл ещё раз.")
        return

    template = TEMPLATES.get(template_id)
    if not template:
        await query.edit_message_text("❌ Неизвестный шаблон.")
        return

    audio_info = pending_audio.pop(user_id)

    if template_id == "spec":
        pending_spec_context[user_id] = {
            "mode": "initial",
            "path": audio_info["path"],
            "chat_id": query.message.chat_id,
        }
        await query.edit_message_text(
            "📐 Тех. задание\n\n"
            "Опиши задачу — что за продукт, для кого, какие цели?\n"
            "Отправь текстом или голосовым сообщением.\n\n"
            "Или нажми «Пропустить» для стандартного оформления.",
            reply_markup=InlineKeyboardMarkup([[
                InlineKeyboardButton("Пропустить →", callback_data=f"spec_skip:{user_id}")
            ]]),
        )
        return

    await _run_standard_template(query, context.bot, user_id, template_id, audio_info["path"], os.path.splitext(audio_info.get("filename", ""))[0])


async def _run_standard_template(query, bot, user_id: str, template_id: str, tmp_path: str, source_filename: str = ""):
    template = TEMPLATES[template_id]
    loop = asyncio.get_event_loop()
    try:
        await query.edit_message_text(f"✅ Шаблон: {template['name']}\n⏳ Извлекаю аудио…")
        try:
            compressed_path = await loop.run_in_executor(None, extract_audio, tmp_path)
        except Exception:
            compressed_path = tmp_path
        file_size_mb = os.path.getsize(compressed_path) / 1024 / 1024
        est_min = max(1, int(file_size_mb / 2))
        await query.edit_message_text(
            f"✅ Шаблон: {template['name']}\n⏳ Транскрибирую… Примерное время: {est_min}–{est_min*2} мин. Ожидайте."
        )
        transcript = await loop.run_in_executor(None, transcribe_audio, compressed_path)
        if compressed_path != tmp_path:
            try:
                os.unlink(compressed_path)
            except OSError:
                pass

        if not transcript.strip():
            await query.edit_message_text("❌ Deepgram не смог распознать речь. Попробуй снова.")
            return

        await query.edit_message_text(
            f"✅ Транскрипция готова ({len(transcript)} символов)\n⏳ Форматирую…"
        )

        tasks = []
        if ANTHROPIC_API_KEY:
            try:
                formatted = await loop.run_in_executor(None, format_with_claude, template_id, transcript)
                tasks = getattr(format_with_claude, '_last_tasks', [])
            except Exception as e:
                logger.exception("Claude failed")
                await bot.send_message(chat_id=query.message.chat_id, text=f"⚠️ Claude ошибка: {e}\n\nСырой транскрипт:")
                formatted = transcript
        else:
            formatted = transcript

        await query.edit_message_text(f"✅ Готово! Шаблон: {template['name']}")
        await deliver_result(query.message.chat_id, bot, user_id, template_id, transcript, formatted, source_filename, tasks)

    except Exception as exc:
        logger.exception("Error processing audio")
        await query.edit_message_text(f"❌ Ошибка: {exc}")
    finally:
        try:
            os.unlink(tmp_path)
        except OSError:
            pass


async def handle_spec_skip(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Skip spec context — process with standard prompt."""
    query = update.callback_query
    await query.answer()
    user_id = query.data.split(":", 1)[1]

    if user_id not in pending_spec_context:
        await query.edit_message_text("❌ Сессия устарела. Отправь аудио заново.")
        return

    spec_state = pending_spec_context.pop(user_id)
    await query.edit_message_text("⏳ Обрабатываю запись…")
    await process_after_spec_context(
        chat_id=query.message.chat_id,
        status_message=query.message,
        bot=context.bot,
        user_id=user_id,
        spec_state=spec_state,
        custom_context=None,
    )


async def handle_action(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handle action buttons: reformat, docx, email, cancels."""
    query = update.callback_query
    await query.answer()

    parts = query.data.split(":", 2)
    action = parts[1]
    user_id = parts[2] if len(parts) == 3 else None

    # ── Cancels ──
    if action == "reformat_cancel":
        await query.edit_message_text("Отменено.")
        return

    if not user_id or user_id not in last_result:
        await query.answer("Результат устарел. Отправь аудио заново.", show_alert=True)
        return

    result = last_result[user_id]

    # ── Reformat ──
    if action == "reformat":
        await query.edit_message_text(
            "Выбери новый шаблон:",
            reply_markup=build_template_keyboard(user_id, prefix="retpl"),
        )

    # ── DOCX ── (кнопки остаются — можно жать несколько действий подряд)
    elif action == "docx":
        docx_path = None
        try:
            src_fn = result.get("source_filename", "")
            docx_path = create_docx(result["template_name"], result["formatted"], src_fn)
            docx_name = f"{src_fn} — {result['template_name']}.docx" if src_fn else f"{result['template_name']}.docx"
            with open(docx_path, "rb") as f:
                await context.bot.send_document(
                    chat_id=query.message.chat_id,
                    document=f,
                    filename=docx_name,
                )
        except Exception as exc:
            logger.exception("DOCX generation failed")
            await query.message.reply_text(f"❌ Ошибка генерации DOCX: {exc}")
        finally:
            if docx_path:
                try:
                    os.unlink(docx_path)
                except OSError:
                    pass

    # ── Markdown ── (кнопки остаются)
    elif action == "md":
        md_path = None
        try:
            src_fn = result.get("source_filename", "")
            md_path = create_md(result["template_id"], result["template_name"], result["formatted"], src_fn)
            md_name = f"{src_fn} — {result['template_name']}.md" if src_fn else f"{result['template_name']}.md"
            with open(md_path, "rb") as f:
                await context.bot.send_document(
                    chat_id=query.message.chat_id,
                    document=f,
                    filename=md_name,
                )
        except Exception as exc:
            logger.exception("MD generation failed")
            await query.message.reply_text(f"❌ Ошибка генерации MD: {exc}")
        finally:
            if md_path:
                try:
                    os.unlink(md_path)
                except OSError:
                    pass


async def handle_reformat_selection(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Re-format stored transcript with a different template."""
    query = update.callback_query
    await query.answer()

    parts = query.data.split(":", 2)
    if parts[0] != "retpl" or len(parts) != 3:
        return
    _, template_id, user_id = parts

    if user_id not in last_result:
        await query.edit_message_text("❌ Результат устарел. Отправь аудио заново.")
        return

    transcript = last_result[user_id]["transcript"]
    template = TEMPLATES.get(template_id)
    if not template:
        await query.edit_message_text("❌ Неизвестный шаблон.")
        return

    if template_id == "spec":
        pending_spec_context[user_id] = {
            "mode": "reformat",
            "transcript": transcript,
            "chat_id": query.message.chat_id,
        }
        await query.edit_message_text(
            "📐 Тех. задание\n\n"
            "Опиши задачу — что за продукт, для кого, какие цели?\n"
            "Отправь текстом или голосовым, или нажми «Пропустить».",
            reply_markup=InlineKeyboardMarkup([[
                InlineKeyboardButton("Пропустить →", callback_data=f"spec_skip:{user_id}")
            ]]),
        )
        return

    await query.edit_message_text(f"⏳ Переформатирую как {template['name']}…")
    loop = asyncio.get_event_loop()
    try:
        if ANTHROPIC_API_KEY:
            formatted = await loop.run_in_executor(None, format_with_claude, template_id, transcript)
        else:
            formatted = transcript
        await query.edit_message_text(f"✅ Переформатировано: {template['name']}")
        await deliver_result(query.message.chat_id, context.bot, user_id, template_id, transcript, formatted)
    except Exception as exc:
        logger.exception("Reformat failed")
        await query.edit_message_text(f"❌ Ошибка: {exc}")


async def handle_second_brain_link(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Кнопки привязки заметки к встрече (шлёт second-brain, решает он же)."""
    query = update.callback_query
    await query.answer()
    if not (SECOND_BRAIN_URL and SECOND_BRAIN_API_KEY):
        await query.edit_message_text("❌ second-brain не настроен.")
        return
    try:
        async with httpx.AsyncClient(timeout=60) as client:
            resp = await client.post(
                f"{SECOND_BRAIN_URL}/link-action",
                json={"data": query.data},
                headers={"X-API-Key": SECOND_BRAIN_API_KEY},
            )
            resp.raise_for_status()
            result = resp.json()
    except Exception as exc:
        logger.exception("second-brain link-action failed")
        await query.edit_message_text(f"❌ Ошибка привязки: {exc}")
        return
    markup = None
    if result.get("button"):
        b = result["button"]
        markup = InlineKeyboardMarkup([[
            InlineKeyboardButton(b["text"], callback_data=b["callback_data"])
        ]])
    await query.edit_message_text(result.get("text", "Готово."), reply_markup=markup)


# ─── Applaud webhook (Plaud-девайс, без Telegram) ──────────────────────────────


def _verify_applaud_signature(raw_body: bytes, header: str) -> bool:
    if not APPLAUD_WEBHOOK_SECRET:
        return True  # секрет не настроен в applaud — как и он сам, не подписываем
    expected = "sha256=" + hmac.new(
        APPLAUD_WEBHOOK_SECRET.encode(), raw_body, hashlib.sha256
    ).hexdigest()
    return hmac.compare_digest(header or "", expected)


async def handle_applaud_webhook(request: web.Request) -> web.Response:
    raw_body = await request.read()
    if not _verify_applaud_signature(raw_body, request.headers.get("X-Applaud-Signature", "")):
        logger.warning("applaud webhook: bad signature")
        return web.Response(status=401, text="bad signature")

    payload = json.loads(raw_body)
    event = payload.get("event")
    if event != "transcript_ready":
        return web.Response(status=200, text="ignored")

    content = payload.get("content") or {}
    transcript = (content.get("transcript_text") or "").strip()
    recording = payload.get("recording") or {}
    if not transcript:
        logger.info("applaud webhook: transcript_ready без текста, пропуск (%s)", recording.get("id"))
        return web.Response(status=200, text="empty transcript")

    if not MY_CHAT_ID:
        logger.error("applaud webhook: MY_CHAT_ID не задан, некуда доставить результат")
        return web.Response(status=500, text="MY_CHAT_ID not configured")

    start_ms = recording.get("start_time_ms")
    start_dt = datetime.fromtimestamp(start_ms / 1000, tz=MSK) if start_ms else datetime.now(MSK)
    title = recording.get("filename") or "Plaud recording"
    source_filename = f"{start_dt:%Y-%m-%d_%H_%M_%S} {title}"

    bot = request.app["bot"]
    chat_id = int(MY_CHAT_ID)
    logger.info("applaud webhook: transcript_ready «%s» (%d символов)", title, len(transcript))

    async def process():
        try:
            loop = asyncio.get_event_loop()
            formatted = await loop.run_in_executor(None, format_with_claude, "meeting", transcript)
            tasks = getattr(format_with_claude, "_last_tasks", [])
        except Exception:
            logger.exception("applaud webhook: Claude formatting failed")
            formatted = transcript
            tasks = []
        await deliver_result(chat_id, bot, "applaud", "meeting", transcript, formatted, source_filename, tasks)

    asyncio.create_task(process())
    return web.Response(status=200, text="accepted")


async def start_webhook_server(bot) -> web.AppRunner:
    webapp = web.Application()
    webapp["bot"] = bot
    webapp.router.add_post("/webhook/applaud", handle_applaud_webhook)
    runner = web.AppRunner(webapp)
    await runner.setup()
    site = web.TCPSite(runner, "127.0.0.1", APPLAUD_WEBHOOK_PORT)
    await site.start()
    logger.info("Applaud webhook listening on 127.0.0.1:%d/webhook/applaud", APPLAUD_WEBHOOK_PORT)
    return runner


# ─── Entry point ──────────────────────────────────────────────────────────────


async def run():
    application = Application.builder().token(TELEGRAM_BOT_TOKEN).build()

    application.add_handler(CommandHandler("start", cmd_start))
    application.add_handler(MessageHandler(
        filters.AUDIO | filters.VOICE | filters.VIDEO | filters.Document.AUDIO | filters.Document.VIDEO,
        handle_audio,
    ))
    application.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_text_message))
    application.add_handler(MessageHandler(
        filters.Document.FileExtension("txt") | filters.Document.FileExtension("md"),
        handle_transcript_file,
    ))
    application.add_handler(CallbackQueryHandler(handle_template_selection, pattern=r"^tpl:"))
    application.add_handler(CallbackQueryHandler(handle_reformat_selection, pattern=r"^retpl:"))
    application.add_handler(CallbackQueryHandler(handle_spec_skip, pattern=r"^spec_skip:"))
    application.add_handler(CallbackQueryHandler(handle_action, pattern=r"^action:"))
    application.add_handler(CallbackQueryHandler(todoist_callback, pattern=r"^todoist:"))
    application.add_handler(CallbackQueryHandler(handle_second_brain_link, pattern=r"^sb(lk|mv):"))

    await application.initialize()
    await application.start()
    await application.updater.start_polling(
        drop_pending_updates=True,
        allowed_updates=[Update.MESSAGE, Update.CALLBACK_QUERY],
    )

    webhook_runner = await start_webhook_server(application.bot)

    logger.info("Bot is running…")
    stop_event = asyncio.Event()
    loop = asyncio.get_running_loop()
    try:
        import signal
        for sig in (signal.SIGTERM, signal.SIGINT):
            loop.add_signal_handler(sig, stop_event.set)
    except (NotImplementedError, ImportError):
        pass

    await stop_event.wait()

    await webhook_runner.cleanup()
    await application.updater.stop()
    await application.stop()
    await application.shutdown()


if __name__ == "__main__":
    asyncio.run(run())
